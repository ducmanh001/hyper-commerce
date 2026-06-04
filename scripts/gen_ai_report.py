#!/usr/bin/env python3
"""Generate HyperCommerce AI Workflow Report (DOCX)"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.0)

# ── Style helpers ─────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0,70,127)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
    return p

def para(text="", bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def table_row_shade(table, row_idx, rgb_hex="DCE6F1"):
    tr = table.rows[row_idx]._tr
    trPr = tr.get_or_add_trPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    trPr.append(shd)

def add_table(headers, rows, col_widths=None, header_bg="1F4E79"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255,255,255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_bg)
        cell._tc.get_or_add_tcPr().append(shd)
    # data rows
    for ri, row in enumerate(rows):
        shade = "F2F7FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
        table_row_shade(t, ri + 1, shade)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ═══════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HyperCommerce")
set_font(run, size=28, bold=True, color=(0,70,127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI Development Workflow Model")
set_font(run, size=20, bold=True, color=(31,78,121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("4-Layer Context Architecture — Token Budget Analysis — Optimization Report")
set_font(run, size=12, color=(89,89,89))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Ngày: {datetime.date.today().strftime('%d/%m/%Y')}  |  Version: 1.0  |  Model: Claude Sonnet 4.6")
set_font(run, size=11, color=(89,89,89))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. TỔNG QUAN MÔ HÌNH
# ═══════════════════════════════════════════════════════════════
heading("1. Tổng Quan Mô Hình AI Workflow", level=1)

para("HyperCommerce sử dụng mô hình 4-Layer Context Cache để tối ưu token khi làm việc với GitHub Copilot (Claude Sonnet 4.6). Thay vì load toàn bộ context mỗi request, context được chia thành 4 tầng với trigger khác nhau:", size=11)

doc.add_paragraph()
add_table(
    ["Layer", "File", "Trigger", "Mục đích", "Tokens (~)"],
    [
        ["L1 — Root", "copilot-instructions.md", "Luôn luôn", "Stack, port, patterns, security", "~1,015t"],
        ["L2 — Service", "agents/{domain}.agent.md", "applyTo: service folder path", "Domain entities, service rules", "~1,000–1,600t"],
        ["L3 — Module", "instructions/*.instructions.md", "applyTo: file type/path", "Code conventions, DB rules", "~453–562t mỗi file"],
        ["L4 — Task", "prompts/*.prompt.md", "Explicit /command", "Spec + checklist từng bước", "~1,400t (on-demand)"],
    ],
    col_widths=[3.0, 4.5, 4.0, 5.0, 2.5]
)

doc.add_paragraph()
para("Nguyên tắc cốt lõi:", bold=True, size=11)
bullet("Layer 1 luôn load → chỉ chứa thông tin vĩnh viễn cần thiết mọi request")
bullet("Layer 2 load theo service đang mở → domain knowledge cụ thể")
bullet("Layer 3 load theo file type → conventions & patterns (backend ≠ frontend)")
bullet("Layer 4 chỉ load khi user invoke /command → spec + checklist task cụ thể")
bullet("On-demand references (SCHEMA.md, EVENTS.md) → load khi cần, không auto-load")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. AUDIT: ĐIỂM YẾU & NHỮNG GÌ ĐÃ FIX
# ═══════════════════════════════════════════════════════════════
heading("2. Audit — Điểm Yếu & Kết Quả Fix", level=1)

heading("2.1 Danh sách vấn đề phát hiện", level=2)

add_table(
    ["#", "Vấn đề", "File bị ảnh hưởng", "Tác động (token)", "Trạng thái"],
    [
        ["1", "nestjs.instructions.md applyTo 'apps/**' → load cho apps/web/ (frontend)\nBackend rules tải cho Next.js = sai & lãng phí", "nestjs.instructions.md", "−562t mỗi web edit", "✅ ĐÃ FIX"],
        ["2", "security.instructions.md cùng applyTo 'apps/**' như nestjs\n2 file cùng scope = overlap không cần thiết cho web", "security.instructions.md", "−453t mỗi web edit", "✅ ĐÃ FIX"],
        ["3", "nextjs.instructions.md không có security section\nSau khi security.instructions.md bị exclude khỏi web, frontend mất hết security rules", "nextjs.instructions.md", "Giảm security risk", "✅ ĐÃ FIX"],
        ["4", "Không có Kafka Event catalog tập trung\nMỗi agent file tự list events riêng → trùng lặp, dễ sai", "Tất cả agent files", "Giảm ~200t, tăng accuracy", "✅ TẠO EVENTS.md"],
        ["5", "copilot-instructions.md không document 4-layer model\nCopilot không biết context hierarchy → load thừa hoặc hỏi lại", "copilot-instructions.md", "Tăng accuracy", "✅ ĐÃ THÊM"],
        ["6", "infra.agent.md có 'Makefile Targets' section\nNội dung đã có trong Makefile → thông tin trùng lặp", "infra.agent.md", "−200t", "✅ ĐÃ XÓA"],
        ["7", "architect.agent.md có BFF section (NOT implemented)\n& Versioning Strategy (generic) → context không có giá trị", "architect.agent.md", "−400t", "✅ ĐÃ XÓA"],
        ["8", "database.instructions.md không reference SCHEMA.md\nAI viết migration không biết table nào đã tồn tại", "database.instructions.md", "Giảm SQL bug risk", "✅ ĐÃ FIX"],
        ["9", "add-feature.prompt.md không có spec input & self-verify\nAI generate code mà không có spec → thiếu context, không verify", "add-feature.prompt.md", "Tăng code quality", "✅ REWRITE"],
    ],
    col_widths=[0.6, 6.5, 4.0, 2.5, 2.0]
)

doc.add_paragraph()
heading("2.2 Điểm yếu còn lại (chưa fix)", level=2)

add_table(
    ["#", "Vấn đề", "File", "Priority", "Đề xuất fix"],
    [
        ["1", "ai-ml.agent.md (6,408 bytes) — 8 code blocks TypeScript dài\nEmbedding pipeline, fraud detection, feed ranking code",
         "ai-ml.agent.md", "HIGH", "Di chuyển code examples vào instructions/ml-patterns.instructions.md. Giữ agent ở mức architecture overview."],
        ["2", "backend.agent.md — 10 code blocks, nhiều nhất toàn bộ hệ thống\nKafka/Redis/BullMQ code trùng với nestjs.instructions.md",
         "backend.agent.md", "HIGH", "Remove code blocks trùng lặp với nestjs.instructions.md. Giảm từ 4,945 → ~2,500 bytes."],
        ["3", "Không có libs/grpc/PROTOS.md catalog\nKhi AI viết gRPC code không biết services/methods available",
         "libs/grpc/src/proto/", "MEDIUM", "Tạo PROTOS.md tương tự EVENTS.md với danh sách service, method, request/response types."],
        ["4", "feature-dev.chatmode.md rất mỏng (25 lines)\nKhông hướng dẫn AI load domain agent phù hợp",
         "feature-dev.chatmode.md", "MEDIUM", "Thêm bước 'Identify domain from active file path → load matching agent' trước khi bắt đầu."],
        ["5", "write-tests.prompt.md thiếu test execution step\nGenerate tests nhưng không chạy test để verify",
         "write-tests.prompt.md", "LOW", "Thêm step cuối: npm run test -- {file} --coverage và minimum threshold check."],
        ["6", "Chưa có instructions/ml-patterns.instructions.md\nCode AI viết cho ai-service/search-service không có convention file",
         "Missing file", "MEDIUM", "Tạo file này với: OpenAI API patterns, Qdrant upsert pattern, embedding cache pattern, batch processing."],
    ],
    col_widths=[0.6, 6.0, 3.5, 1.8, 4.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. TOKEN BUDGET — BEFORE vs AFTER
# ═══════════════════════════════════════════════════════════════
heading("3. Token Budget — Before vs After Optimization", level=1)

para("Công thức tính: 1 token ≈ 3.8 ký tự (English/code mix). File sizes lấy từ wc -c.", size=11)
doc.add_paragraph()

heading("3.1 Edit file frontend (apps/web/src/app/page.tsx)", level=2)

add_table(
    ["File", "Bytes", "Tokens", "Trước fix", "Sau fix"],
    [
        ["copilot-instructions.md (L1)", "3,858", "~1,015t", "✅ Load", "✅ Load"],
        ["frontend.agent.md (L2)", "3,370", "~887t", "✅ Load", "✅ Load"],
        ["nestjs.instructions.md (L3)", "2,135", "~562t", "❌ Load (WRONG)", "⛔ Không load"],
        ["nextjs.instructions.md (L3)", "2,763", "~727t", "✅ Load", "✅ Load (+ security)"],
        ["security.instructions.md (L3)", "1,722", "~453t", "❌ Load (WRONG)", "⛔ Không load"],
        ["TỔNG AUTO-LOADED", "", "", "~3,644t", "~2,629t"],
        ["Tiết kiệm", "", "", "—", "−1,015t (−28%)"],
    ],
    col_widths=[5.5, 2.0, 2.0, 3.0, 3.0]
)

doc.add_paragraph()
heading("3.2 Edit file backend (apps/order-service/src/order.service.ts)", level=2)

add_table(
    ["File", "Bytes", "Tokens", "Trước fix", "Sau fix"],
    [
        ["copilot-instructions.md (L1)", "3,858", "~1,015t", "✅ Load", "✅ Load"],
        ["commerce.agent.md (L2)", "3,935", "~1,035t", "✅ Load", "✅ Load"],
        ["nestjs.instructions.md (L3)", "2,135", "~562t", "✅ Load", "✅ Load"],
        ["security.instructions.md (L3)", "1,722", "~453t", "✅ Load", "✅ Load"],
        ["TỔNG AUTO-LOADED", "", "", "~3,065t", "~3,065t"],
        ["Ghi chú", "", "", "—", "Backend không đổi, correct từ đầu"],
    ],
    col_widths=[5.5, 2.0, 2.0, 3.0, 3.0]
)

doc.add_paragraph()
heading("3.3 Edit migration (infrastructure/postgres/migrations/005.sql)", level=2)

add_table(
    ["File", "Bytes", "Tokens", "Load?"],
    [
        ["copilot-instructions.md (L1)", "3,858", "~1,015t", "✅ Luôn load"],
        ["infra.agent.md (L2)", "5,139", "~1,352t", "✅ Load (applyTo: infrastructure/**)"],
        ["database.instructions.md (L3)", "2,222", "~585t", "✅ Load (applyTo: infrastructure/**)"],
        ["SCHEMA.md (on-demand ref)", "6,728", "~1,771t", "📎 Load khi AI cần biết table hiện tại"],
        ["TỔNG", "", "", "~2,952t (auto) + ~1,771t (on-demand)"],
    ],
    col_widths=[5.5, 2.0, 2.0, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. CHI PHÍ THEO LOẠI TASK
# ═══════════════════════════════════════════════════════════════
heading("4. Chi Phí Token Theo Loại Task", level=1)

para("Công thức: Total = Auto-loaded context + Layer 4 prompt + User request + AI response", size=11)
para("Giá Claude Sonnet 4.6 (Copilot): ~$3/1M input tokens + ~$15/1M output tokens", size=11, color=(89,89,89))
doc.add_paragraph()

add_table(
    ["Loại Task", "Ví dụ cụ thể", "Context\n(auto)", "Prompt\n(L4)", "AI Response\n(output)", "Tổng tokens", "Chi phí ước tính"],
    [
        ["Tiny fix",
         "Fix typo trong DTO field name",
         "~3,065t", "—", "~200t", "~3,365t", "~$0.004"],
        ["Small feature",
         "Add product wishlist (1 entity, 3 endpoints)",
         "~3,085t", "~1,397t", "~3,500t", "~8,082t", "~$0.062"],
        ["Medium feature",
         "Review & Rating system (3 entities, 5 endpoints, 1 saga)",
         "~3,085t", "~1,397t", "~7,000t", "~13,482t", "~$0.114"],
        ["Large feature",
         "Flash Sale với saga + fraud check + live integration",
         "~3,065t", "~1,397t +\n~2,947t refs", "~12,000t", "~21,409t", "~$0.213"],
        ["New service",
         "/add-service: chat-service từ đầu",
         "~3,065t", "~2,662t\n(add-service)", "~8,000t", "~15,727t", "~$0.142"],
        ["Migration write",
         "Tạo migration cho feature mới (5 tables)",
         "~2,952t\n+SCHEMA", "~1,397t", "~2,000t", "~8,120t", "~$0.066"],
        ["Code review",
         "/code-review toàn bộ PR (10 files)",
         "~3,065t", "—", "~5,000t", "~8,165t", "~$0.081"],
        ["Debug session",
         "Debug Kafka consumer không consume",
         "~3,065t", "—", "~1,500t", "~4,665t", "~$0.030"],
    ],
    col_widths=[2.8, 4.5, 2.0, 2.0, 2.5, 2.2, 2.5]
)

doc.add_paragraph()
para("* Chi phí tính theo giá Copilot Enterprise/Business. Với Copilot Individual, token không tính thêm phí riêng — ảnh hưởng chủ yếu là tốc độ phản hồi và độ chính xác.", size=10, color=(89,89,89))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. FILE INVENTORY — CURRENT STATE
# ═══════════════════════════════════════════════════════════════
heading("5. Inventory Toàn Bộ Context Files", level=1)

heading("5.1 Agent Files (Layer 2)", level=2)

add_table(
    ["File", "Lines", "Bytes", "applyTo (scope)", "Code Blocks", "Đánh giá"],
    [
        ["orchestrator.agent.md", "31", "1,765", "None (manual-load)", "0", "✅ Gọn, đúng vai trò routing"],
        ["frontend.agent.md", "97", "3,370", "apps/web/**", "~3", "✅ Tốt, không có code ví dụ thừa"],
        ["commerce.agent.md", "109", "3,935", "apps/{order,payment,inventory,review}-service/**", "~4", "✅ Tốt"],
        ["social.agent.md", "119", "4,011", "apps/{user,feed,live,subscription,chat}-service/**", "~4", "✅ Tốt"],
        ["platform.agent.md", "126", "4,692", "apps/{notification,analytics,admin,ads}-service/**", "~5", "✅ Tốt"],
        ["backend.agent.md", "151", "4,945", "libs/**", "~10", "⚠️ Quá nhiều code blocks, trùng nestjs.instructions"],
        ["infra.agent.md", "155", "5,139", "infrastructure/**,docker-compose.yml,Dockerfile*,Makefile", "~8", "⚠️ Còn khá nặng sau trim"],
        ["architect.agent.md", "145", "5,189", "apps/api-gateway/**", "~6", "✅ OK sau trim, cần verify code blocks"],
        ["ai-ml.agent.md", "166", "6,408", "apps/{ai,search,analytics}-service/**,libs/algorithms/**", "~8", "❌ Nặng nhất, cần split"],
    ],
    col_widths=[3.8, 1.2, 1.2, 5.0, 2.0, 3.5]
)

doc.add_paragraph()
heading("5.2 Instruction Files (Layer 3)", level=2)

add_table(
    ["File", "Lines", "Bytes", "applyTo (scope)", "Trạng thái"],
    [
        ["security.instructions.md", "44", "1,722", "apps/*-service/**/*.ts,libs/**/*.ts", "✅ Đã fix applyTo"],
        ["nestjs.instructions.md", "60", "2,135", "apps/*-service/**/*.ts,libs/**/*.ts", "✅ Đã fix applyTo"],
        ["database.instructions.md", "54", "2,222", "infrastructure/**,apps/*/src/entities/**", "✅ Đã thêm SCHEMA.md reference"],
        ["testing.instructions.md", "59", "1,891", "apps/**/*.spec.ts,*.e2e-spec.ts,*.test.ts", "✅ Tốt"],
        ["nextjs.instructions.md", "61", "2,763", "apps/web/**/*.tsx,apps/web/**/*.ts", "✅ Đã thêm security section"],
        ["ml-patterns.instructions.md", "—", "—", "MISSING", "❌ Cần tạo cho ai-service/search-service"],
    ],
    col_widths=[4.5, 1.2, 1.2, 5.5, 4.5]
)

doc.add_paragraph()
heading("5.3 Prompt Files (Layer 4)", level=2)

add_table(
    ["File", "Lines", "Bytes", "Có self-verify?", "Trạng thái"],
    [
        ["write-tests.prompt.md", "56", "2,173", "⚠️ Partial (1 ref)", "Thiếu: test execution step"],
        ["add-service.prompt.md", "74", "2,662", "✅ Có (3 refs)", "Tốt"],
        ["refactor.prompt.md", "75", "2,618", "✅ Có (3 refs)", "Tốt"],
        ["delete-feature.prompt.md", "77", "2,973", "⚠️ Partial (2 refs)", "Thiếu: FK orphan + Kafka consumer cleanup check"],
        ["migrate-service.prompt.md", "107", "3,940", "✅ Có (5 refs)", "Tốt"],
        ["add-feature.prompt.md", "138", "5,308", "✅ Mạnh (6 refs + full checklist)", "✅ Tốt nhất — spec + 9-step + self-verify"],
    ],
    col_widths=[4.5, 1.2, 1.2, 3.0, 6.5]
)

doc.add_paragraph()
heading("5.4 On-Demand Reference Files", level=2)

add_table(
    ["File", "Lines", "Bytes", "Mục đích", "Load khi nào"],
    [
        ["libs/events/EVENTS.md", "55", "4,471", "Kafka event catalog: 20 events, saga flow, DLT topics", "User/AI hỏi về cross-service events"],
        ["infrastructure/postgres/SCHEMA.md", "~200", "6,728", "25 bảng DB, indexes, FKs, next migration number", "Viết migration / tạo entity mới"],
        ["libs/grpc/PROTOS.md", "—", "—", "gRPC services & methods catalog", "❌ MISSING — Cần tạo"],
    ],
    col_widths=[4.5, 1.5, 1.5, 6.0, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. VÍ DỤ THỰC TẾ: ADD WISHLIST FEATURE
# ═══════════════════════════════════════════════════════════════
heading("6. Ví Dụ Thực Tế: Implement 'Product Wishlist'", level=1)

para("Dưới đây là flow hoàn chỉnh khi developer dùng /add-feature để thêm tính năng Wishlist vào user-service:", size=11)
doc.add_paragraph()

steps = [
    ("Bước 1", "Developer mở file apps/user-service/src/app.module.ts",
     "Copilot tự động load:\n• L1: copilot-instructions.md (~1,015t)\n• L2: social.agent.md (~1,055t) — vì applyTo: apps/user-service/**\n• L3: nestjs.instructions.md (~562t)\n• L3: security.instructions.md (~453t)\nTotal auto-context: ~3,085t"),
    ("Bước 2", "Developer gõ /add-feature",
     "Copilot load thêm:\n• L4: add-feature.prompt.md (~1,397t)\nTotal context: ~4,482t"),
    ("Bước 3", "Developer điền spec vào Step 0:\nFeature: product-wishlist | Service: user-service:3001\nNew table: wishlists (id, user_id, product_id, created_at)\nAPI: POST /api/v1/wishlists, DELETE /api/v1/wishlists/:id\nEvents: wishlist.added → feed-service (tracking)\nUI: /account/wishlist page",
     "AI có đủ context để generate:\n• wishlists.entity.ts\n• add-to-wishlist.dto.ts\n• wishlist.service.ts\n• wishlist.controller.ts\n• 005_wishlist_table.sql migration\n• API gateway route\n• /account/wishlist/page.tsx"),
    ("Bước 4", "AI generate code (output ~3,500t)",
     "AI output includes:\n• Entity với @Entity, @PrimaryGeneratedColumn('uuid'), userId shard key\n• Migration đúng số thứ tự (check SCHEMA.md)\n• DTO với class-validator decorators\n• Service với ownership check\n• Controller với @UseGuards(JwtAuthGuard)\n• Frontend page với Server Component pattern"),
    ("Bước 5", "Self-verify (Step 9 trong prompt)",
     "AI tự kiểm tra:\n☑ Entity trong TypeOrmModule.forFeature([])\n☑ Controller trong module.controllers[]\n☑ Migration file có số đúng\n☑ SCHEMA.md được update\n☑ API gateway route added\n☑ API client throws on non-ok (no mock)"),
    ("Bước 6", "Chạy make verify",
     "Automated check:\n1/4 TypeScript — no errors\n2/4 ESLint — passed\n3/4 Security scan — no Math.random(), no MOCK_ imports\n4/4 Wiring — entity in forFeature, migration sequence ok\n→ ✓ verify passed"),
]

for step_name, action, result in steps:
    p = doc.add_paragraph()
    run = p.add_run(f"► {step_name}: {action}")
    set_font(run, size=11, bold=True, color=(31,78,121))

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(result)
    set_font(run, size=10.5, color=(50,50,50))
    doc.add_paragraph()

add_table(
    ["Metric", "Giá trị"],
    [
        ["Tổng input tokens", "~4,482t (context + prompt)"],
        ["Tổng output tokens", "~3,500t (generated code)"],
        ["Total tokens per request", "~7,982t"],
        ["Chi phí ước tính", "~$0.066"],
        ["Thời gian (ước tính)", "~15-25 giây (Claude Sonnet 4.6)"],
        ["Số file được generate", "7 files"],
        ["Tỷ lệ code đúng lần đầu (estimate)", "~85% (nhờ spec + self-verify)"],
    ],
    col_widths=[6.0, 10.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. ROADMAP TỐI ƯU TIẾP THEO
# ═══════════════════════════════════════════════════════════════
heading("7. Roadmap Tối Ưu Tiếp Theo", level=1)

add_table(
    ["Priority", "Task", "File cần thay đổi", "Impact", "Token saving"],
    [
        ["P1 — HIGH",
         "Trim backend.agent.md\nXóa code blocks trùng với nestjs.instructions.md\nGiữ architecture overview + module dependency map",
         "backend.agent.md", "Giảm context nặng nhất cho libs/**",
         "~1,000t mỗi libs edit"],
        ["P1 — HIGH",
         "Split ai-ml.agent.md\nAgent: architecture overview + model versions\nMới: instructions/ml-patterns.instructions.md với code examples",
         "ai-ml.agent.md +\nnew instruction file",
         "Giảm agent size, tăng reusability",
         "~800t mỗi ai-service edit"],
        ["P2 — MEDIUM",
         "Tạo libs/grpc/PROTOS.md\nCatalog: service name, methods, request/response types\nTương tự EVENTS.md cho Kafka",
         "New: libs/grpc/PROTOS.md",
         "Giảm gRPC code bug",
         "On-demand, ~500t khi cần"],
        ["P2 — MEDIUM",
         "Nâng cấp feature-dev.chatmode.md\nThêm domain auto-detection step\nLink đến EVENTS.md + SCHEMA.md khi cần",
         "feature-dev.chatmode.md",
         "AI load đúng agent hơn",
         "Tăng accuracy"],
        ["P3 — LOW",
         "Fix write-tests.prompt.md\nThêm step: npm run test -- {file} --coverage\nKiểm tra minimum 80% coverage",
         "write-tests.prompt.md",
         "Tests không chỉ generate mà còn pass",
         "Tăng quality"],
        ["P3 — LOW",
         "Fix delete-feature.prompt.md\nThêm FK orphan check + Kafka consumer cleanup\nDestructive operation cần verify mạnh hơn",
         "delete-feature.prompt.md",
         "Giảm nguy cơ regression khi xóa feature",
         "Giảm bug risk"],
    ],
    col_widths=[2.0, 5.5, 4.0, 3.0, 2.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. KIẾN TRÚC CONTEXT FLOW (TEXT DIAGRAM)
# ═══════════════════════════════════════════════════════════════
heading("8. Context Flow Diagram", level=1)

diagram = """
Developer mở file  →  GitHub Copilot nhận active file path
         │
         ▼
┌─────────────────────────────────────────────────────────────┐
│  LAYER 1 — Root (LUÔN LOAD)                                 │
│  copilot-instructions.md                                    │
│  Stack · Port map · 4 patterns · Security rules · Routing  │
│  ~1,015 tokens                                              │
└─────────────────────────────────────────────────────────────┘
         │
         ▼  applyTo matching by file path
┌─────────────────────────────────────────────────────────────┐
│  LAYER 2 — Service Agent (1 file, path-triggered)           │
│  social.agent.md      (user/feed/live service)              │
│  commerce.agent.md    (order/payment/inventory/review)      │
│  frontend.agent.md    (apps/web/**)                         │
│  platform.agent.md    (notification/analytics/ads/admin)    │
│  ai-ml.agent.md       (ai/search/analytics service)         │
│  backend.agent.md     (libs/**)                             │
│  infra.agent.md       (infrastructure/**)                   │
│  architect.agent.md   (apps/api-gateway/**)                 │
│  ~880–1,600 tokens                                          │
└─────────────────────────────────────────────────────────────┘
         │
         ▼  applyTo matching by file extension/type
┌─────────────────────────────────────────────────────────────┐
│  LAYER 3 — Module Instructions (1-2 files)                  │
│  Backend file (.ts in *-service or libs):                   │
│    nestjs.instructions.md + security.instructions.md        │
│  Frontend file (.tsx/.ts in apps/web):                      │
│    nextjs.instructions.md (includes security rules)         │
│  DB/Migration file:                                         │
│    database.instructions.md                                 │
│  Test file:                                                 │
│    testing.instructions.md                                  │
│  ~450–1,000 tokens                                          │
└─────────────────────────────────────────────────────────────┘
         │
         ▼  user explicit invoke: /add-feature, /refactor, etc.
┌─────────────────────────────────────────────────────────────┐
│  LAYER 4 — Task Prompt (on-demand)                          │
│  add-feature.prompt.md   /add-feature                       │
│  add-service.prompt.md   /add-service                       │
│  write-tests.prompt.md   /write-tests                       │
│  refactor.prompt.md      /refactor                          │
│  delete-feature.prompt.md /delete-feature                   │
│  migrate-service.prompt.md /migrate-service                 │
│  ~1,400–2,700 tokens                                        │
└─────────────────────────────────────────────────────────────┘
         │
         ▼  AI references when needed (NOT auto-loaded)
┌─────────────────────────────────────────────────────────────┐
│  ON-DEMAND REFERENCES                                       │
│  libs/events/EVENTS.md        ← 20 Kafka events catalog    │
│  infrastructure/postgres/SCHEMA.md  ← 25 tables, next migration │
│  libs/grpc/PROTOS.md (TODO)   ← gRPC services catalog      │
└─────────────────────────────────────────────────────────────┘

TOTAL PER REQUEST: ~2,350–4,500t auto + up to 4,000t on-demand
"""

p = doc.add_paragraph()
run = p.add_run(diagram)
run.font.name = "Courier New"
run.font.size = Pt(8.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. KẾT LUẬN
# ═══════════════════════════════════════════════════════════════
heading("9. Kết Luận", level=1)

para("Mô hình 4-Layer Context Cache đã được triển khai đầy đủ cho HyperCommerce với những cải tiến chính:", size=11)
doc.add_paragraph()

bullet("Layer 1 (Root): Slim, compact — chỉ chứa thông tin cần thiết mọi request (~1,015t)")
bullet("Layer 2 (Service): 8 domain agents, tự động load theo file path — không cần user hướng dẫn")
bullet("Layer 3 (Module): 5 instruction files với applyTo chính xác — backend ≠ frontend, tránh overlap")
bullet("Layer 4 (Task): 6 prompt files với structured spec input + self-verify checklist")
bullet("On-demand refs: EVENTS.md + SCHEMA.md — 4,471 + 6,728 bytes chỉ load khi AI cần")
bullet("make verify: 4-check automated validator (TypeScript + ESLint + Security + Wiring)")

doc.add_paragraph()
para("Điểm yếu còn lại cần xử lý theo priority:", bold=True, size=11)
bullet("P1: Trim backend.agent.md (10 code blocks) + split ai-ml.agent.md (~−1,800t tổng)")
bullet("P2: Tạo libs/grpc/PROTOS.md catalog + nâng cấp feature-dev chatmode")
bullet("P3: Fix write-tests.prompt.md + delete-feature.prompt.md")

doc.add_paragraph()
para("Token budget tổng kết:", bold=True, size=11)
add_table(
    ["Loại task", "Context load", "Output estimate", "Total", "Chi phí"],
    [
        ["Tiny fix", "~3,065t", "~200t", "~3,365t", "~$0.004"],
        ["Small feature", "~4,482t", "~3,500t", "~8,082t", "~$0.062"],
        ["Large feature", "~7,409t", "~12,000t", "~21,409t", "~$0.213"],
        ["New service", "~5,727t", "~8,000t", "~15,727t", "~$0.142"],
    ],
    col_widths=[3.5, 3.0, 3.5, 2.5, 2.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Được tạo tự động bởi GitHub Copilot (Claude Sonnet 4.6) — HyperCommerce AI Workflow Report")
set_font(run, size=9, color=(150,150,150))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────────
out = "LCB_TOT_Framework_AI_Agent.docx"
doc.save(out)
print(f"✅ Saved: {out}")
